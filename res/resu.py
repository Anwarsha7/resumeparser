import streamlit as st
import spacy
import fitz  # PyMuPDF for PDFs
import docx
import re

# Load spaCy NLP model
nlp = spacy.load("en_core_web_sm")

# Function to extract text from PDF including tables

def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = "\n".join(page.get_text("text") for page in doc)
    return text

# Function to extract text and tables from DOCX
def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    
    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(" | ".join(row_data))
        tables.append("\n".join(table_data))
    
    return text + "\n\nTables:\n" + "\n\n".join(tables) if tables else text

# Function to extract key details using NLP
def extract_resume_details(text):
    doc = nlp(text)

    # Extract names (first proper noun in the text)
    name = None
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text
            break

    # Extract email
    email = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    email = email[0] if email else None

    # Extract phone numbers (basic regex pattern)
    phone = re.findall(r"\+?\d{10,15}", text)
    phone = phone[0] if phone else None

    # Extract skills (based on simple keyword matching)
    skills = []
    skill_keywords = {"ann"}
    for token in doc:
        if token.text in skill_keywords:
            skills.append(token.text)

    return {
        "Name": name,
        "Email": email,
        "Phone": phone,
        "Skills": ", ".join(set(skills)) if skills else "Not found"
    }

# Streamlit Frontend
st.markdown("<h1 style='color: white;'>📄 Resume Parser using NLP</h1>", unsafe_allow_html=True)
st.markdown("<h3 style='color: #1E90FF;'>Parsed Information</h3>", unsafe_allow_html=True)

st.write("Upload a *PDF* or *DOCX* resume, and the app will extract key details using NLP, including tables.")

uploaded_file = st.file_uploader("Upload Resume", type=["pdf", "docx"])

if uploaded_file is not None:
    st.success("File uploaded successfully!")

    if uploaded_file.type == "application/pdf":
        extracted_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        extracted_text = extract_text_from_docx(uploaded_file)
    
    if extracted_text:
        st.subheader("Extracted Text and Tables")
        st.text_area("Resume Content", extracted_text, height=300)

        st.subheader("Parsed Information")
        parsed_data = extract_resume_details(extracted_text)
        
        for key, value in parsed_data.items():
            st.write(f"*{key}:* {value}")
    else:
        st.error("Could not extract text from the uploaded file.")