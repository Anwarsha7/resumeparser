# Flask & Web
from flask import Flask, render_template, request, redirect, session, send_from_directory, flash, url_for
from flask_mail import Mail, Message

# Database
from pymongo import MongoClient
from bson.objectid import ObjectId

# Security & Auth
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

# File Processing
import pdfplumber
import docx
import PyPDF2
from pdfminer.high_level import extract_text as pdfminer_extract
from pdfminer.layout import LAParams

# NLP
import spacy
from spacy.matcher import Matcher, PhraseMatcher

# Utilities
import os
from datetime import datetime, timedelta
from dotenv import load_dotenv
import traceback
import re
import phonenumbers
from phonenumbers import carrier, geocoder
from math import ceil
from concurrent.futures import ThreadPoolExecutor
import concurrent.futures
import csv
from io import StringIO

# Timezone Handling
from pytz import timezone
import tzlocal

# Fuzzy Matching
from fuzzywuzzy import fuzz

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY')

# Configure Flask-Mail
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT'))
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'False').lower() in ['true', '1', 'yes']
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')

mail = Mail(app)

# MongoDB Atlas connection
mongo_uri = os.getenv('MONGO_URI')
try:
    client = MongoClient(mongo_uri)
    db = client['resume_parser']
    admins_collection = db['admins']
    users_collection = db['users']
    resumes_collection = db['resumes']
    print("Connected to MongoDB Atlas successfully!")
except Exception as e:
    print(f"Error connecting to MongoDB Atlas: {e}")

# Mail
from flask import request, redirect, url_for, flash
from flask_mail import Mail, Message

# Ensure Mail is initialized somewhere in your app setup:
# mail = Mail(app)

@app.route('/contact', methods=['POST'])
def handle_contact():
    if request.method == 'POST':
        name = request.form.get('name')
        email = request.form.get('email')
        subject = request.form.get('subject')  # <-- get subject
        message = request.form.get('message')

        if not all([name, email, message]):
            flash('Please fill all fields', 'error')
            return redirect(url_for('contact'))

        try:
            msg = Message(
                subject=f"New Contact Form: {subject or 'No Subject'} from {name}",
                sender=app.config['MAIL_USERNAME'],
                recipients=[app.config['MAIL_USERNAME']],
                body=f"""Name: {name}
Email: {email}
Subject: {subject or 'No Subject'}

Message:
{message}"""
            )
            msg.reply_to = email
            mail.send(msg)
            flash('Your message has been sent successfully!', 'success')
        except Exception as e:
            print(f"Failed to send email: {str(e)}")
            flash('Failed to send your message. Please try again later.', 'error')

        return redirect(url_for('contact'))





# Timezone configuration
app.config['TIMEZONE'] = os.getenv('APP_TIMEZONE', 'UTC')
print(f"Using Timezone: {app.config['TIMEZONE']}")

def get_local_time():
    """Gets the current time in the server's local timezone"""
    return datetime.now()

# Upload folder configuration
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Load spaCy model for English only
try:
    nlp = spacy.load("en_core_web_sm")  # Using small English model for basic NER
except OSError:
    print("ERROR: spaCy English model 'en_core_web_sm' not found.")
    print("Please run: python -m spacy download en_core_web_sm")
    exit()

# ================== IMPROVED SKILL EXTRACTION ================== #
 
# Skill database (can be loaded from CSV or API)
SKILLS_DB = {
    # Programming Languages
    "python": "Programming",
    "java": "Programming",
    "javascript": "Programming",
    "c++": "Programming",
    "c#": "Programming",
    "ruby": "Programming",
    "php": "Programming",
    "swift": "Programming",
    "kotlin": "Programming",
    "go": "Programming",
    "typescript": "Programming",
    
    # Web Development
    "html": "Web",
    "css": "Web",
    "react": "Web",
    "angular": "Web",
    "vue": "Web",
    "node.js": "Web",
    "django": "Web",
    "flask": "Web",
    "spring": "Web",
    
    # Databases
    "mysql": "Database",
    "postgresql": "Database",
    "mongodb": "Database",
    "sqlite": "Database",
    "oracle": "Database",
    
    # Cloud & DevOps
    "aws": "Cloud",
    "azure": "Cloud",
    "google cloud": "Cloud",
    "docker": "DevOps",
    "kubernetes": "DevOps",
    "terraform": "DevOps",
    
    # AI/ML
    "tensorflow": "AI/ML",
    "pytorch": "AI/ML",
    "scikit-learn": "AI/ML",
    "opencv": "AI/ML",
    
    # Other
    "git": "Tools",
    "linux": "OS",
    "windows": "OS"
}

# Skill synonyms
SKILL_SYNONYMS = {
    "js": "javascript",
    "aws cloud": "amazon web services",
    "tf": "tensorflow",
    "postgres": "postgresql",
    "gcp": "google cloud"
}

def normalize_skill_name(skill):
    skill_lower = skill.lower()
    return SKILL_SYNONYMS.get(skill_lower, skill_lower)

def build_skill_matcher(nlp):
    matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
    patterns = [nlp(skill) for skill in SKILLS_DB.keys()]
    matcher.add("SKILLS", patterns)
    return matcher

skill_matcher = build_skill_matcher(nlp)

def extract_skills_with_context(text):
    doc = nlp(text)
    matches = skill_matcher(doc)
    skills = []

    for match_id, start, end in matches:
        skill_span = doc[start:end]
        skill_text = skill_span.text
        
        # Check for experience level (e.g., "Python (5 years)")
        exp_pattern = r"\((\d+)\s+(?:year|yr)s?\)"
        exp_match = re.search(exp_pattern, skill_text)
        exp_level = exp_match.group(1) if exp_match else None

        # Check for proficiency (e.g., "Advanced Python")
        proficiency = None
        if start > 0 and doc[start-1].text.lower() in ["expert", "advanced", "intermediate", "beginner"]:
            proficiency = doc[start-1].text

        normalized_skill = normalize_skill_name(skill_text)
        category = SKILLS_DB.get(normalized_skill, "Other")

        skill = {
            "name": skill_text,
            "normalized_name": normalized_skill,
            "category": category,
            "experience_years": exp_level,
            "proficiency": proficiency
        }
        skills.append(skill)

    return skills

def extract_skills_hybrid(text):
    """Returns only skill names (strings) without any metadata"""
    # Get all skills with context (internal processing)
    skills_with_metadata = _extract_skills_with_metadata(text)
    
    # Extract just the names
    skill_names = [skill['name'] for skill in skills_with_metadata]
    
    # Add regex fallback matches
    for skill in SKILLS_DB:
        if re.search(rf"\b{re.escape(skill)}\b", text, re.IGNORECASE):
            skill_names.append(skill)
    
    # Return sorted, unique list
    return sorted(list(set(skill_names)))

def _extract_skills_with_metadata(text):
    """Internal helper that does the actual NLP processing"""
    doc = nlp(text)
    matches = skill_matcher(doc)
    skills = []
    
    for match_id, start, end in matches:
        skill_span = doc[start:end]
        skills.append({
            'name': skill_span.text,  # We'll only use this externally
            # Internal metadata (not returned)
            'normalized_name': normalize_skill_name(skill_span.text),
            'category': SKILLS_DB.get(normalize_skill_name(skill_span.text), "Other")
        })
    return skills
 
# ================== END OF IMPROVED SKILL EXTRACTION ================== #

# Enhanced education patterns
DEGREES = [
    "bachelor", "bsc", "b.tech", "be", "bs", "ba", "b.com", "bca", "b.arch", "b.pharm", "b.ed", "bba",
    "master", "msc", "m.tech", "me", "ms", "ma", "m.com", "mca", "m.arch", "m.pharm", "m.ed", "mba",
    "phd", "doctorate"
]

# ================== UPDATED TEXT EXTRACTION FUNCTIONS ================== #
def extract_text_from_pdf(filepath):
    """Enhanced PDF text extraction with dedicated table handling"""
    text = ""
    tables_text = ""
    
    # First try pdfplumber with table extraction
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                # Extract regular text with layout awareness
                page_text = page.extract_text(
                    layout=True,
                    x_tolerance=2,
                    y_tolerance=2,
                    keep_blank_chars=False
                ) or ""
                
                # Extract tables separately
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        tables_text += "\n[TABLE START]\n"
                        for row in table:
                            # Clean each cell and handle None values
                            clean_row = [str(cell).strip() if cell is not None else "" for cell in row]
                            tables_text += "| " + " | ".join(clean_row) + " |\n"
                        tables_text += "[TABLE END]\n"
                
                text += page_text + "\n"
                
    except Exception as e:
        print(f"pdfplumber error: {e}. Trying fallback methods.")
        text = ""
    
    # Fallback to PyPDF2 if pdfplumber fails
    if len(text.strip()) < 100:
        try:
            with open(filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                if reader.is_encrypted:
                    try:
                        reader.decrypt('')
                    except:
                        pass
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
        except Exception as e:
            print(f"PyPDF2 error: {e}")
            text = ""
    
    # Final fallback to pdfminer
    if len(text.strip()) < 100:
        try:
            laparams = LAParams(
                line_margin=0.5,
                char_margin=2.0,
                word_margin=0.1
            )
            text = pdfminer_extract(filepath, laparams=laparams)
        except Exception as e:
            print(f"pdfminer error: {e}")
            text = ""
    
    # Combine text and tables
    full_text = text + tables_text
    
    # Advanced cleanup
    full_text = re.sub(r'\s+', ' ', full_text).strip()
    full_text = re.sub(r'[^\x00-\x7F]+', ' ', full_text)  # Remove non-ASCII chars
    full_text = re.sub(r'\s*[\n\r]+\s*', '\n', full_text)  # Normalize newlines
    
    return full_text

def extract_text_from_docx(filepath):
    """Enhanced DOCX text extraction with improved table handling"""
    full_text = []
    
    try:
        doc = docx.Document(filepath)
        
        # Improved paragraph extraction
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:  # Only add non-empty paragraphs
                full_text.append(para_text)
        
        # Enhanced table extraction with Markdown format
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Extract text from all paragraphs in cell
                    cell_text = ' '.join([p.text for p in cell.paragraphs]).strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    table_data.append(row_data)
            
            # Convert table to Markdown format if we have data
            if table_data:
                # Create Markdown table header
                md_table = []
                if len(table_data) > 0:
                    md_table.append('| ' + ' | '.join(table_data[0]) + ' |')
                    md_table.append('| ' + ' | '.join(['---'] * len(table_data[0])) + ' |')
                
                # Add table rows
                for row in table_data[1:]:
                    md_table.append('| ' + ' | '.join(row) + ' |')
                
                full_text.append('\n'.join(md_table))
        
        # Extract headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                header_text = header.text.strip()
                if header_text:
                    full_text.append(header_text)
            for footer in section.footer.paragraphs:
                footer_text = footer.text.strip()
                if footer_text:
                    full_text.append(footer_text)
    
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
    
    return "\n".join(full_text)
# ================== END OF UPDATED TEXT EXTRACTION FUNCTIONS ================== #


# Enhanced entity extraction with improved accuracy
def extract_entities(text):
    """Enhanced entity extraction with improved accuracy and skill detection"""
    # Pre-process text with better normalization
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII chars
    
    # Name extraction with improved patterns
    name = "Unknown"
    name_patterns = [
        r'^(?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',  # First line name
        r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b',  # Anywhere in text
        r'(?:Resume|CV|Curriculum Vitae)\s+of\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        r'(?:Name|Full Name|Candidate)[:\s]*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        r'Personal\s+Details[^a-z]*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)'
    ]
    
    for pattern in name_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            name = match.group(1) if len(match.groups()) > 0 else match.group(0)
            name = ' '.join([part for part in name.split() if len(part) > 1])  # Filter single letters
            break
    
    # Fallback to spaCy NER with better filtering
    if name == "Unknown":
        doc = nlp(text[:2000])  # Check more text for name
        for ent in doc.ents:
            if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
                name = ent.text
                break
    
    # Email extraction with stricter validation
    email = "Unknown"
    email_matches = re.finditer(
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b', 
        text
    )
    for match in email_matches:
        if len(match.group(0).split('@')[0]) > 1:  # Minimum username length
            email = match.group(0)
            break
    
    # Phone - improved international pattern
    phone = "Unknown"
    phone_matches = re.finditer(
        r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', 
        text
    )
    for match in phone_matches:
        phone = match.group(0)
        if len(phone) > 7:  # Basic validation
            break
    
    # Use improved skill extraction
    skills = extract_skills_hybrid(text)
    
    # Enhanced Education Extraction
    education = []
    edu_sections = re.finditer(
        r'(?i)(?:education|academic\s+(?:background|qualifications?)|degrees?|qualifications?)[:\s]*(.*?)(?=\n\s*\n|\Z)', 
        text,
        re.DOTALL
    )
    
    for section in edu_sections:
        edu_text = section.group(1)
        if not edu_text:
            continue
            
        # Split education items by bullets or newlines
        edu_items = re.split(r'\n\s*[•\-*]\s*|\n\s+', edu_text)
        
        for item in edu_items:
            item = item.strip()
            if not item or len(item) < 10:  # Skip very short items
                continue
                
            # Extract degree type
            degree_match = re.search(
                r'(?i)\b(' + '|'.join(DEGREES) + r')\b', 
                item
            )
            
            if degree_match:
                degree = degree_match.group(1)
                
                # Extract institution with better patterns
                institution_match = re.search(
                    r'(?:at\s+|,\s+|-\s+|from\s+|@\s+|in\s+)?([A-Z][a-zA-Z\s&]+(?:University|College|Institute|School|Academy|Polytechnic)\b|[A-Z][a-zA-Z\s&]+)', 
                    item[degree_match.end():],
                    re.IGNORECASE
                )
                
                institution = institution_match.group(1).strip() if institution_match else "Unknown Institution"
                
                # Extract year with better patterns
                year_match = re.search(
                    r'(?:(?:19|20)\d{2}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+(?:19|20)\d{2})', 
                    item,
                    re.IGNORECASE
                )
                year = year_match.group(0) if year_match else "Year not specified"
                
                education.append(f"{degree} from {institution} ({year})")
    
    # Enhanced Experience Calculation
    exp_years = 0
    
    # First look for explicit experience mentions
    exp_matches = re.finditer(
        r'(?i)(\d+)\s+(?:year|yr)s?\s+(?:experience|exp\.?|of\s+experience)', 
        text
    )
    for match in exp_matches:
        years = int(match.group(1))
        if years > exp_years:
            exp_years = years
    
    # Then look for date ranges in employment history
    if exp_years == 0:
        date_ranges = re.finditer(
            r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}\s*[-–—]\s*(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}|\d{4}\s*[-–—]\s*\d{4}|\bpresent\b|\bcurrent\b',
            text,
            re.IGNORECASE
        )
        
        current_year = datetime.now().year
        total_months = 0
        
        for dr in date_ranges:
            dates = re.findall(r'(?:19|20)\d{2}', dr.group(0))
            if len(dates) == 2:
                start_year, end_year = map(int, dates)
                total_months += (end_year - start_year) * 12
            elif 'present' in dr.group(0).lower() or 'current' in dr.group(0).lower():
                if dates:  # Has start year
                    start_year = int(dates[0])
                    total_months += (current_year - start_year) * 12
        
        if total_months > 0:
            exp_years = round(total_months / 12)
    
    # Enhanced Location Extraction
    location = "Unknown"
    
    # First try to find address pattern
    address_match = re.search(
        r'\b\d+\s+[\w\s]+,\s*(?:[A-Z]{2,}\s*\d{5,6})?\s*[\w\s]+(?:,\s*[A-Z]{2,})?', 
        text
    )
    if address_match:
        location = address_match.group(0)
    else:
        # Then look for city/country patterns
        location_match = re.search(
            r'\b(?:lives? in|located in|based in|address|from)\s*[:]?\s*([A-Z][a-zA-Z\s]+(?:City|Town|Village)?)', 
            text, 
            re.IGNORECASE
        )
        if location_match:
            location = location_match.group(1).strip()
        else:
            # Fallback to spaCy
            doc = nlp(text[:3000])  # Process more text for location
            for ent in doc.ents:
                if ent.label_ in ("GPE", "LOC") and len(ent.text) > 2:
                    location = ent.text
                    break
    
    # Enhanced Certifications Extraction
    certifications = []
    cert_patterns = [
        r'\b(?:AWS|Azure|Google Cloud|GCP|CCNA|CCNP|CCIE|CEH|CISSP|PMP|CAPM|CSM|CSPO|ITIL|COBIT|TOGAF|OCP|OCA|MCSA|MCSE|MCSD|RHCE|RHCSA|VCP|VCAP|VCDX|CKA|CKAD)\b',
        r'\b(?:Certified|Professional|Associate|Specialist|Expert|Master)\b[\w\s]+?(?:Certification|Certificate|Certified)\b',
        r'\b[A-Z]{2,}(?:\s+[A-Z]{2,})*\s+(?:Certification|Certificate)\b',
        r'\b(?:Microsoft|Oracle|Cisco|Amazon|Google|IBM)\s+[A-Za-z]+\s+(?:Certification|Certificate)\b'
    ]
    
    for pattern in cert_patterns:
        cert_matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in cert_matches:
            cert = match.group(0).strip()
            
            # Validate certificate format
            if (len(cert.split()) >= 2 or  # At least two words
                any(acronym in cert for acronym in ['AWS', 'CCNA', 'PMP', 'CISSP', 'ITIL']) or
                re.search(r'[A-Z]{3,}', cert)):  # Contains acronyms
                
                # Extract issuing organization if available
                org_match = re.search(
                    r'(?:by\s+|from\s+|@\s+|\(|\s)([A-Z][a-zA-Z\s&]+(?:Inc\.?|Corp\.?|LLC|Ltd\.?)?)', 
                    text[match.end():match.end()+150]
                )
                
                if org_match:
                    cert += f" ({org_match.group(1).strip()})"
                
                if cert not in certifications:
                    certifications.append(cert)
    
    return {
        'name': name.title() if name != "Unknown" else name,
        'email': email.lower() if email != "Unknown" else email,
        'phone': phone,
        'skills': skills,
        'education': education,
        'experience': exp_years,
        'location': location,
        'certifications': certifications
    }

# Main function to parse a resume file with duplicate prevention
def parse_resume(file_path):
    """Main function to parse a resume file with improved error handling"""
    filename = os.path.basename(file_path)
    
    try:
        # First check if this file was already parsed
        existing = resumes_collection.find_one({'filename': filename, 'parsed': True})
        if existing:
            print(f"Resume {filename} was already parsed")
            return existing

        if filename.lower().endswith('.pdf'):
            extracted_text = extract_text_from_pdf(file_path)
        elif filename.lower().endswith('.docx'):
            extracted_text = extract_text_from_docx(file_path)
        elif filename.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
        else:
            return {
                'filename': filename,
                'uploaded_at': get_local_time(),
                'parsed': False,
                'parsing_error': 'Unsupported file type',
                'raw_text': ''
            }

        if not extracted_text or len(extracted_text.strip()) < 30:
            return {
                'filename': filename,
                'uploaded_at': get_local_time(),
                'parsed': False,
                'parsing_error': 'Text extraction failed or text too short',
                'raw_text': extracted_text[:500]
            }

        entities = extract_entities(extracted_text)

        result = {
            'name': entities['name'],
            'email': entities['email'],
            'phone': entities['phone'],
            'skills': entities['skills'],
            'education': entities['education'],
            'experience': entities['experience'],
            'location': entities['location'],
            'certifications': entities['certifications'],
            'filename': filename,
            'uploaded_at': get_local_time(),
            'parsed': True,
            'parsed_at': get_local_time(),
            'parsing_error': None,
            'raw_text': extracted_text[:1000]  # Store first 1000 chars for reference
        }

        return result

    except Exception as e:
        return {
            'filename': filename,
            'uploaded_at': get_local_time(),
            'parsed': False,
            'parsing_error': f"Error: {str(e)}",
            'raw_text': ''
        }

# Function to save parsed resume to database with duplicate prevention
def save_parsed_resume(data):
    # Check if resume with same filename already exists
    existing = resumes_collection.find_one({'filename': data['filename'], 'parsed': True})
    if existing:
        print(f"Resume {data['filename']} already exists in database")
        return existing['_id']

    # Set default values
    data.setdefault('parsed', False)
    data.setdefault('parsing_error', 'Unknown error during save' if not data.get('parsed') else None)
    data.setdefault('name', 'Unknown')
    data.setdefault('email', 'Unknown')
    data.setdefault('phone', 'Unknown')
    data.setdefault('skills', [])
    data.setdefault('education', [])
    data.setdefault('experience', 0)
    data.setdefault('location', 'Unknown')
    data.setdefault('certifications', [])
    data.setdefault('filename', 'unknown_filename')
    data.setdefault('uploaded_at', get_local_time())
    data.setdefault('parsed_at', get_local_time() if data['parsed'] else None)
    data.setdefault('raw_text', '')

    try:
        result = resumes_collection.insert_one(data)
        return result.inserted_id
    except Exception as e:
        print(f"Database save error: {e}")
        return None

# Function to get parsed resumes
def get_parsed_resumes():
    return list(resumes_collection.find({'parsed': True}).sort('uploaded_at', -1))

# Calculate match percentage (unchanged)
def calculate_match_percentage(resume, request_args):
    """Calculate how well the resume matches filter criteria"""
    match_details = {
        'skills_match': 0,
        'experience_match': 0,
        'education_match': 0,
        'location_match': 0
    }
    
    total_weight = 0
    total_score = 0
    
    # Skills matching (weight: 40%)
    if 'skills' in request_args and request_args['skills']:
        required_skills = [s.strip().lower() for s in request_args['skills'].split(',')]
        resume_skills = [s.lower() for s in resume.get('skills', [])]
        
        if required_skills and resume_skills:
            matched = len(set(required_skills) & set(resume_skills))
            match_details['skills_match'] = (matched / len(required_skills)) * 100
            total_score += match_details['skills_match'] * 0.4
            total_weight += 0.4
    
    # Experience matching (weight: 30%)
    if 'min_experience' in request_args and request_args['min_experience']:
        min_exp = int(request_args['min_experience'])
        resume_exp = resume.get('experience', 0)
        
        if resume_exp >= min_exp:
            match_details['experience_match'] = 100
        else:
            match_details['experience_match'] = (resume_exp / min_exp) * 100
        
        total_score += match_details['experience_match'] * 0.3
        total_weight += 0.3
    
    # Education matching (weight: 20%)
    if 'education' in request_args and request_args['education']:
        required_edu = request_args['education'].lower()
        resume_edu = ' '.join(resume.get('education', [])).lower()
        
        if required_edu in resume_edu:
            match_details['education_match'] = 100
        
        total_score += match_details['education_match'] * 0.2
        total_weight += 0.2
    
    # Location matching (weight: 10%)
    if 'location' in request_args and request_args['location']:
        required_loc = request_args['location'].lower()
        resume_loc = resume.get('location', '').lower()
        
        if required_loc in resume_loc:
            match_details['location_match'] = 100
        
        total_score += match_details['location_match'] * 0.1
        total_weight += 0.1
    
    # Calculate overall match percentage
    if total_weight > 0:
        overall_match = total_score / total_weight
    else:
        overall_match = 0
    
    # Add match details to resume for template access
    resume['match_details'] = match_details
    
    return round(overall_match)

# Context processor (unchanged)
@app.context_processor
def utility_processor():
    return dict(calculate_match_percentage=calculate_match_percentage)

# Bulk delete route (unchanged)
@app.route('/delete-selected-resumes', methods=['POST'])
def delete_selected_resumes():
    if not session.get('admin_logged_in'):
        flash('Please log in to perform this action.', 'error')
        return redirect('/adminlogin')

    resume_ids = request.form.getlist('resume_ids')
    referrer = request.form.get('referrer', url_for('dashboard'))

    if not resume_ids:
        flash('No resumes selected for deletion.', 'error')
        return redirect(referrer)

    try:
        # Convert string IDs to ObjectId
        object_ids = [ObjectId(rid) for rid in resume_ids]
        
        # Find and delete resumes
        deleted_count = 0
        for resume in resumes_collection.find({'_id': {'$in': object_ids}}):
            try:
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], resume['filename'])
                if os.path.exists(file_path):
                    os.remove(file_path)
                deleted_count += 1
            except Exception as e:
                print(f"Error deleting file for resume {resume['_id']}: {e}")
                continue
        
        # Delete from database
        result = resumes_collection.delete_many({'_id': {'$in': object_ids}})
        
        flash(f'Successfully deleted {deleted_count} resume(s)!', 'success')
    except Exception as e:
        flash(f'Error deleting resumes: {str(e)}', 'error')

    return redirect(referrer)

# Updated upload route with duplicate prevention and multi-file support
@app.route('/upload', methods=['GET', 'POST'])
def upload_resume():
    if not session.get('admin_logged_in'):
        flash('Please log in to upload a resume.', 'error')
        return redirect('/adminlogin')

    if request.method == 'POST':
        name = request.form.get('name')
        email = request.form.get('email')
        files = request.files.getlist('resume')

        if not files or all(file.filename == '' for file in files):
            flash('No files selected for upload.', 'error')
            return redirect(url_for('upload_resume'))

        uploaded_count = 0
        errors = []
        processed_files = []

        def process_file(file):
            try:
                if file.filename == '':
                    return {'success': False, 'message': 'Empty filename'}

                if not allowed_file(file.filename):
                    return {'success': False, 'message': f'Invalid file type for {file.filename}'}

                filename = secure_filename(file.filename)
                
                # Check if file already exists in database (only check unparsed files)
                existing_unparsed = resumes_collection.find_one({
                    'filename': filename,
                    'parsed': False
                })
                
                # Check if parsed version exists
                existing_parsed = resumes_collection.find_one({
                    'filename': filename,
                    'parsed': True
                })
                
                if existing_parsed:
                    return {'success': False, 'message': f'File {filename} already exists and was parsed'}
                
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                
                # If there's an unparsed version, delete it first
                if existing_unparsed:
                    resumes_collection.delete_one({'_id': existing_unparsed['_id']})
                
                parsed_data = parse_resume(filepath)
                if not parsed_data.get('parsed', False):
                    os.remove(filepath)  # Clean up if parsing failed
                    return {'success': False, 'message': f'Failed to parse {filename}'}

                # Use form data only if no name/email was extracted
                if parsed_data['name'] == 'Unknown':
                    parsed_data['name'] = name
                if parsed_data['email'] == 'Unknown':
                    parsed_data['email'] = email

                parsed_data.update({
                    'parsed': True,
                    'uploaded_at': get_local_time(),
                    'parsed_at': get_local_time()
                })

                save_parsed_resume(parsed_data)
                return {'success': True, 'filename': filename}
                
            except Exception as e:
                # Clean up if any error occurred
                if 'filepath' in locals() and os.path.exists(filepath):
                    os.remove(filepath)
                return {'success': False, 'message': f'Error processing {file.filename}: {str(e)}'}

        # Process files in parallel (up to 4 at a time)
        with ThreadPoolExecutor(max_workers=4) as executor:
            futures = [executor.submit(process_file, file) for file in files]
            
            for future in concurrent.futures.as_completed(futures):
                result = future.result()
                if result['success']:
                    uploaded_count += 1
                    processed_files.append(result['filename'])
                else:
                    errors.append(result['message'])

        if uploaded_count > 0:
            success_msg = f'Successfully processed {uploaded_count} resume(s)'
            if uploaded_count <= 5:
                success_msg += ': ' + ', '.join(processed_files)
            flash(success_msg, 'success')
        
        if errors:
            error_msg = f'Encountered {len(errors)} error(s)'
            if len(errors) <= 5:
                error_msg += ': ' + ', '.join(errors)
            flash(error_msg, 'error')

        return redirect(url_for('parsed_resumes'))

    return render_template('upload.html')

# ===== MODIFIED FILTER ROUTES =====

@app.route('/advanced-filter', methods=['GET'])
def advanced_filter():
    if not session.get('admin_logged_in'):
        flash('Please log in to filter resumes.', 'error')
        return redirect('/adminlogin')

    # Get all filter parameters
    skills = request.args.get('skills')
    min_experience = request.args.get('min_experience')
    education = request.args.get('education')
    location = request.args.get('location')
    certification = request.args.get('certification')
    date_range = request.args.get('date_range')
    sort_by = request.args.get('sort_by', 'uploaded_at')
    sort_order = int(request.args.get('sort_order', -1))

    # Build the query based on filters
    query = {'parsed': True}

    # Skills filter (multiple skills can be selected)
    if skills:
        skills_list = [skill.strip().lower() for skill in skills.split(',')]
        query['skills'] = {'$all': [{'$elemMatch': {'$regex': skill, '$options': 'i'}} for skill in skills_list]}

    # Minimum experience filter
    if min_experience:
        query['experience'] = {'$gte': int(min_experience)}

    # Education level filter
    if education:
        query['education'] = {'$regex': education, '$options': 'i'}

    # Location filter
    if location:
        query['location'] = {'$regex': location, '$options': 'i'}

    # Certification filter
    if certification:
        query['certifications'] = {'$regex': certification, '$options': 'i'}

    # Date range filter
    if date_range:
        now = datetime.utcnow()
        if date_range == 'today':
            start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
        elif date_range == 'week':
            start_date = now - timedelta(days=now.weekday())
            start_date = start_date.replace(hour=0, minute=0, second=0, microsecond=0)
        elif date_range == 'month':
            start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        query['uploaded_at'] = {'$gte': start_date}

    # Validate sort field
    valid_sort_fields = ['name', 'email', 'experience', 'education', 'uploaded_at']
    if sort_by not in valid_sort_fields:
        sort_by = 'uploaded_at'

    # Get filtered and sorted resumes
    filtered_resumes = list(resumes_collection.find(query).sort(sort_by, sort_order))

    return render_template('parsed_resumes.html', 
                         resumes=filtered_resumes,
                         skills=skills,
                         min_experience=min_experience,
                         education=education,
                         location=location,
                         certification=certification,
                         date_range=date_range,
                         sort_by=sort_by,
                         sort_order=sort_order)

# Simplified filter route that uses advanced_filter
@app.route('/filter', methods=['GET'])
def filter():
    return advanced_filter()

# Simplified sort route that uses advanced_filter
@app.route('/sort', methods=['GET'])
def sort():
    return advanced_filter()

# Main parsed resumes route
@app.route('/parsed-resumes')
def parsed_resumes():
    if not session.get('admin_logged_in'):
        flash('Please log in to view parsed resumes.', 'error')
        return redirect('/adminlogin')

    # Check if we have any filter parameters
    has_filters = any(param in request.args for param in ['skills', 'min_experience', 'education', 
                                                         'location', 'certification', 'date_range'])
    
    # If we have filters, use the advanced filter
    if has_filters:
        return advanced_filter()
    
    # Default case - no filters, just show all parsed resumes
    resumes = list(resumes_collection.find({'parsed': True}).sort('uploaded_at', -1))

    return render_template('parsed_resumes.html',
                         resumes=resumes)

# Delete Resume Route (unchanged)
@app.route('/delete-resume/<resume_id>', methods=['POST'])
def delete_resume(resume_id):
    if not session.get('admin_logged_in'):
        flash('Only admins can delete resumes.', 'error')
        return redirect(url_for('admin_login'))

    try:
        if ',' in resume_id:
            resume_ids = resume_id.split(',')
            deleted_count = 0
            
            for rid in resume_ids:
                try:
                    resume_oid = ObjectId(rid.strip())
                    resume = resumes_collection.find_one({'_id': resume_oid})
                    
                    if resume:
                        file_path = os.path.join(app.config['UPLOAD_FOLDER'], resume['filename'])
                        if os.path.exists(file_path):
                            try:
                                os.remove(file_path)
                            except Exception as e:
                                print(f"Error deleting file {resume['filename']}: {e}")
                        
                        resumes_collection.delete_one({'_id': resume_oid})
                        deleted_count += 1
                
                except Exception as e:
                    print(f"Error deleting resume {rid}: {e}")
                    continue
            
            flash(f'Successfully deleted {deleted_count} resume(s)!', 'success')
        
        else:
            resume_oid = ObjectId(resume_id)
            resume = resumes_collection.find_one({'_id': resume_oid})

            if resume:
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], resume['filename'])
                if os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except Exception as e:
                        print(f"Error deleting file: {e}")
                        flash('Error deleting file from server.', 'warning')
                
                resumes_collection.delete_one({'_id': resume_oid})
                flash('Resume deleted successfully!', 'success')
            else:
                flash('Resume not found.', 'error')

        referrer = request.form.get('referrer', url_for('dashboard'))
        return redirect(referrer)

    except Exception as e:
        print(f"Error deleting resume: {e}")
        flash(f'Error deleting resume: {str(e)}', 'error')
        return redirect(url_for('dashboard'))

# Download Resume (unchanged)
@app.route('/download/<filename>')
def download_resume(filename):
    if not session.get('admin_logged_in'):
        flash('Please log in to download a resume.', 'error')
        return redirect('/adminlogin')

    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# Route to view a single resume
@app.route('/resume/<resume_id>')
def view_resume(resume_id):
    if not session.get('admin_logged_in'):
        flash('Please log in to view resumes.', 'error')
        return redirect('/adminlogin')

    try:
        resume = resumes_collection.find_one({'_id': ObjectId(resume_id)})
        if not resume:
            flash('Resume not found.', 'error')
            return redirect(url_for('parsed_resumes'))
        
        return render_template('view_resume.html', resume=resume)
    except Exception as e:
        flash(f'Error viewing resume: {str(e)}', 'error')
        return redirect(url_for('parsed_resumes'))

# Route to search resumes
@app.route('/search', methods=['GET'])
def search_resumes():
    if not session.get('admin_logged_in'):
        flash('Please log in to search resumes.', 'error')
        return redirect('/adminlogin')

    query = request.args.get('q', '')

    if not query:
        return redirect(url_for('parsed_resumes'))

    search_query = {
        'parsed': True,
        '$or': [
            {'name': {'$regex': query, '$options': 'i'}},
            {'email': {'$regex': query, '$options': 'i'}},
            {'skills': {'$regex': query, '$options': 'i'}},
            {'education': {'$regex': query, '$options': 'i'}},
            {'location': {'$regex': query, '$options': 'i'}}
        ]
    }

    resumes = list(resumes_collection.find(search_query)
                  .sort('uploaded_at', -1))

    return render_template('parsed_resumes.html',
                         resumes=resumes,
                         search_query=query)

# Route to export resumes
@app.route('/export-resumes', methods=['POST'])
def export_resumes():
    if not session.get('admin_logged_in'):
        flash('Please log in to export resumes.', 'error')
        return redirect('/adminlogin')

    resume_ids = request.form.getlist('resume_ids')
    export_format = request.form.get('export_format', 'json')

    if not resume_ids:
        flash('No resumes selected for export.', 'error')
        return redirect(request.referrer)

    try:
        object_ids = [ObjectId(rid) for rid in resume_ids]
        resumes = list(resumes_collection.find({'_id': {'$in': object_ids}}))

        if export_format == 'json':
            from bson.json_util import dumps
            response = app.response_class(
                response=dumps(resumes),
                status=200,
                mimetype='application/json'
            )
            response.headers['Content-Disposition'] = 'attachment; filename=resumes_export.json'
            return response

        elif export_format == 'csv':
            import csv
            from io import StringIO
            
            si = StringIO()
            cw = csv.writer(si)
            
            # Write header
            cw.writerow(['Name', 'Email', 'Phone', 'Skills', 'Education', 'Experience', 'Location'])
            
            # Write data
            for resume in resumes:
                cw.writerow([
                    resume.get('name', ''),
                    resume.get('email', ''),
                    resume.get('phone', ''),
                    ', '.join(resume.get('skills', [])),
                    ', '.join(resume.get('education', [])),
                    resume.get('experience', 0),
                    resume.get('location', '')
                ])
            
            response = app.response_class(
                response=si.getvalue(),
                status=200,
                mimetype='text/csv'
            )
            response.headers['Content-Disposition'] = 'attachment; filename=resumes_export.csv'
            return response

        else:
            flash('Invalid export format selected.', 'error')
            return redirect(request.referrer)

    except Exception as e:
        flash(f'Error exporting resumes: {str(e)}', 'error')
        return redirect(request.referrer)

# Home Route
@app.route('/')
def index():
    return render_template('front.html')

# About Route
@app.route('/about')
def about():
    return render_template('about.html')

# Contact Route
@app.route('/contact')
def contact():
    return render_template('contact.html')

# Admin Login Route
@app.route('/adminlogin', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        admin = admins_collection.find_one({'username': username})
        
        if admin and check_password_hash(admin['password'], password):
            session['admin_logged_in'] = True
            session['user'] = username
            flash('Login successful!', 'success')
            return redirect('/dashboard')
        
        flash('Invalid credentials, please try again.', 'error')
        return redirect('/adminlogin')

    return render_template('adminlogin.html')

# Dashboard Route
@app.route('/dashboard')
def dashboard():
    if not session.get('admin_logged_in'):
        flash('Please log in to access the dashboard.', 'error')
        return redirect('/adminlogin')

    resumes = list(resumes_collection.find({'parsed': False}).sort('uploaded_at', -1))
    
    for resume in resumes:
        if isinstance(resume['uploaded_at'], str):
            try:
                resume['uploaded_at'] = datetime.strptime(resume['uploaded_at'], 'on %d/%m/%Y at %H:%M:%S')
            except ValueError:
                resume['uploaded_at'] = datetime.utcnow()
    
    return render_template('dashboard.html', resumes=resumes)

# Logout Route
@app.route('/logout')
def logout():
    session.pop('admin_logged_in', None)
    session.pop('user', None)
    flash('You have been logged out.', 'success')
    return redirect('/adminlogin')

# Admin Registration Route
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')

        if admins_collection.find_one({'username': username}):
            flash('Username already exists. Please choose a different username.', 'error')
            return redirect('/register')

        if admins_collection.find_one({'email': email}):
            flash('Email already exists. Please use a different email.', 'error')
            return redirect('/register')

        hashed_password = generate_password_hash(password)

        user_data = {
            'username': username,
            'email': email,
            'password': hashed_password
        }
        admins_collection.insert_one(user_data)

        flash('Registration successful! Please log in.', 'success')
        return redirect('/adminlogin')

    return render_template('register.html')

# User Registration Route
@app.route('/userregister', methods=['GET', 'POST'])
def userregister():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')

        if users_collection.find_one({'username': username}):
            flash('Username already exists. Please choose a different username.', 'error')
            return redirect('/userregister')

        if users_collection.find_one({'email': email}):
            flash('Email already exists. Please use a different email.', 'error')
            return redirect('/userregister')

        hashed_password = generate_password_hash(password)

        user_data = {
            'username': username,
            'email': email,
            'password': hashed_password
        }
        users_collection.insert_one(user_data)

        flash('Registration successful! Please log in.', 'success')
        return redirect('/userlogin')

    return render_template('userregister.html')

# User Login Route
@app.route('/userlogin', methods=['GET', 'POST'])
def user_login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        user = users_collection.find_one({'username': username})

        if user and check_password_hash(user['password'], password):
            session['user_logged_in'] = True
            session['username'] = username
            flash('Login successful!', 'success')
            return redirect('/uploaduser')
        else:
            flash('Invalid credentials, please try again.', 'error')
            return redirect('/userlogin')

    return render_template('userlogin.html')

# User Upload Route
@app.route('/uploaduser', methods=['GET', 'POST'])
def upload_user_resume():
    if not session.get('user_logged_in'):
        flash('Please log in to upload a resume.', 'error')
        return redirect('/userlogin')

    if request.method == 'POST':
        name = request.form.get('name')
        email = request.form.get('email')
        file = request.files.get('resume')

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            # Check if file already exists
            if resumes_collection.find_one({'filename': filename}):
                os.remove(filepath)
                flash('A resume with this filename already exists.', 'error')
                return redirect('/uploaduser')

            user_local_time = request.form.get('local_time')

            resume_data = {
                'filename': filename,
                'name': name,
                'email': email,
                'uploaded_at': user_local_time,
                'parsed': False
            }
            resumes_collection.insert_one(resume_data)

            flash('Uploaded successfully!', 'success')
            return redirect('/uploaduser')
        else:
            flash('Invalid file type. Only PDF, DOCX, and TXT files are allowed.', 'error')

    return render_template('uploaduser.html')

# User Logout Route
@app.route('/userlogout')
def user_logout():
    session.pop('user_logged_in', None)
    session.pop('username', None)
    flash('You have been logged out.', 'success')
    return redirect('/userlogin')

# Error handlers
@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_server_error(e):
    return render_template('500.html'), 500

if __name__ == "__main__":
    app.run(debug=True, use_reloader=False)